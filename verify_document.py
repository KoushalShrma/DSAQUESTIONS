#!/usr/bin/env python3
"""
Verification script to check the DSA questions document structure
"""

from docx import Document
import os

def verify_document():
    doc_path = '/home/runner/work/DSAQUESTIONS/DSAQUESTIONS/DSA_Questions_Collection.docx'
    
    if not os.path.exists(doc_path):
        print("❌ Document not found!")
        return False
    
    try:
        doc = Document(doc_path)
        print("✅ Document loaded successfully!")
        
        # Count paragraphs and headings
        headings = []
        total_paragraphs = len(doc.paragraphs)
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text)
        
        print(f"📊 Document Statistics:")
        print(f"   Total paragraphs: {total_paragraphs}")
        print(f"   Total headings: {len(headings)}")
        
        # Check for main sections
        expected_sections = ['EASY QUESTIONS', 'MEDIUM QUESTIONS', 'HARD QUESTIONS']
        found_sections = []
        
        for heading in headings:
            for section in expected_sections:
                if section in heading:
                    found_sections.append(section)
                    break
        
        print(f"\n📑 Document Sections Found:")
        for section in found_sections:
            print(f"   ✅ {section}")
        
        missing_sections = set(expected_sections) - set(found_sections)
        if missing_sections:
            print(f"\n❌ Missing sections: {missing_sections}")
        else:
            print(f"\n✅ All required sections present!")
        
        # File size check
        file_size = os.path.getsize(doc_path)
        print(f"\n📄 File size: {file_size / 1024:.1f} KB")
        
        if file_size > 40000:  # > 40KB indicates substantial content
            print("✅ Document has substantial content")
        else:
            print("⚠️  Document might be lacking content")
        
        return True
        
    except Exception as e:
        print(f"❌ Error verifying document: {e}")
        return False

if __name__ == "__main__":
    print("🔍 Verifying DSA Questions Document...")
    verify_document()